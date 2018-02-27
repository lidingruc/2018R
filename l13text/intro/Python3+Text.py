
# coding: utf-8

# # Text analysis
# ## 8th Jun 2017
# 
# 
# ## Feng  Li(pkulifeng@163.com)
# 
# ## Based on Tianguang Meng（maxmeng@tsinghua.edu.cn）

# In[ ]:




# ### 1-Input Text

# In[1]:

news = '习近平抵达捷克进行国事访问'


# In[2]:

print (news)


# In[3]:

texts = ['云南去年5800余名公职人员受处分','国务院成立山东疫苗案调查组','东北三省取消玉米临储']


# In[4]:

print (texts[2])


# ## 2-Text Detection

# In[5]:

len("习近平抵达捷克进行国事访问") #length of the string 


# In[6]:

len("Hello World") #length of the string 


# In[7]:

'Alice' == "Alice"


# In[8]:

'Alice' == "alice"


# In[9]:

"That's a nice hat."


# In[10]:

print ("1\t2") #\t is tab


# In[11]:

print ("1\n2") #\n is newline


# In[12]:

print ("云南去年5800余名公职人员受处分\n国务院\t东北三省取消玉米临储") #end-line character


# ## 3-Text Computation

# In[13]:

"White"+" "+"Rabbit"


# In[14]:

"alice".capitalize()


# In[15]:

print ("alice".upper())


# In[16]:

print ("ALIce".lower())


# In[17]:

"There were doors all around the hall".count('e')


# In[18]:

"There were doors all round the hall".count('all')


# In[19]:

"人民币在过去半年总是躺着中枪。英国人完成脱欧，人民币并没有太多暴露于脱欧的风险头寸，然后投资者寻求避险，人民币遭到“错杀”".count('人')


# In[20]:

print ("   White Rabbit   ".strip())
print ("\t White Rabbit  \t ".strip('\t '))


# In[21]:

"Rabbit".startswith("R")


# In[22]:

"Rabbit".endswith("t")


# In[23]:

print ("a" in "Rabbit") #returns a boolean 

print ("doors" in "There were doors all round the hall" )#returns a boolean


# ## 4-Text Data

# In[24]:

id_number={'张三':'54201','李四':'87302','王五':'32201'}


# In[25]:

print (id_number['李四'])


# In[26]:

category={}
category["苹果"]="fruit"
category["香蕉"]="fruit"
category["草莓"]="fruit"
category["土豆"]="vegetable"
category["西红柿"]="vegetable"


# In[27]:

category["香蕉"]


# In[28]:

doc_a = "Brocolli is good to eat. My brother likes to eat good brocolli, but not my mother."
doc_b = "My mother spends a lot of time driving my brother around to baseball practice."
doc_c = "Some health experts suggest that driving may cause increased tension and blood pressure."
doc_d = "I often feel pressure to perform well at school, but my mother never seems to drive my brother to do better."
doc_e = "Health professionals say that brocolli is good for your health."

# compile sample documents into a list
doc_set = [doc_a, doc_b, doc_c, doc_d, doc_e]


# In[29]:

print (doc_set)


# ## 5-Text Reading

# In[39]:

import os
import re
import pandas as pd
import numpy as np


# In[63]:

os.chdir(r"/Users/liding/E/Bdata/Course/6TextasData/")


# In[41]:

report= pd.read_csv('samgov.csv',encoding="gbk")
# 只需要读数据的时候设定编码即可
report.head()
#print report['content'][0].decode('gbk') Python2中需要进行decode显示


# In[42]:

len(report) #length of the string 


# In[62]:

report2 = pd.read_excel('/Users/liding/E/Bdata/Course/6TextasData/sample2.xlsx') # iris数据
report2.head()


# In[67]:

import codecs
text = codecs.open('lianghui17.txt', 'r', encoding='utf-8').read()  #


# In[68]:

print (text)


# In[69]:

len(report2) #length of the string 


# In[74]:

import docx
doc = docx.Document('xi.docx')


# In[75]:

len(doc.paragraphs)
doc.paragraphs[0].text


# ## 6-Text Segment
# 

# In[76]:

import pandas as pd
import numpy as np
import jieba.posseg as pseg
import jieba.analyse
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.naive_bayes import GaussianNB
from sklearn.decomposition import NMF, LatentDirichletAllocation
import jieba
import matplotlib.pyplot as plt
get_ipython().magic('matplotlib inline')

# 清华大学的lac http://thulac.thunlp.org/


# In[45]:

seg_list = jieba.cut("我来到北京清华大学", cut_all=True)
print ("Full Mode:", "/ ".join(seg_list))  # 全模式

seg_list = jieba.cut("我来到北京清华大学", cut_all=False)
print ("Default Mode:", "/ ".join(seg_list))  # 精确模式

seg_list = jieba.cut("他来到了网易杭研大厦")  # 默认是精确模式
print (", ".join(seg_list))

seg_list = jieba.cut_for_search("小明硕士毕业于中国科学院计算所，后在日本京都大学深造")  # 搜索引擎模式
print (", ".join(seg_list))


# In[46]:

words = pseg.cut("我爱北京天安门")
for word, flag in words:
    print('%s %s' % (word, flag)) #词性标注


# In[47]:

result = jieba.tokenize('十动然拒是网络新词')  # 词的起始位置
for tk in result:
    print("word %s\t\t start: %d \t\t end:%d" % (tk[0],tk[1],tk[2]))


# In[48]:

#添加自定义词典
#jieba.load_userdict(filename)
jieba.add_word('十动然拒')
result = jieba.tokenize('十动然拒是网络新词')
for tk in result:
    print("word %s\t\t start: %d \t\t end:%d" % (tk[0],tk[1],tk[2]))


# In[77]:

#可手动调节词典
print('/'.join(jieba.cut('如果放到post中将出错。', HMM=False)))

jieba.suggest_freq(('中', '将'), True)
print('/'.join(jieba.cut('如果放到post中将出错。', HMM=False)))

print('/'.join(jieba.cut('「台中」正确应该不会被切开', HMM=False)))

jieba.suggest_freq('台中', True)
print('/'.join(jieba.cut('「台中」正确应该不会被切开', HMM=False)))


# In[50]:

##基于 TF-idf算法的关键词抽取
#withWeight 为是否一并返回关键词权重值，默认值为 False
s = "高圆圆和赵又廷在京举行答谢宴，诸多明星现身捧场，其中包括张杰、谢娜夫妇、何炅、徐克、张凯丽等。高圆圆身穿粉色外套，看到大批记者在场露出娇羞神色，赵又廷则戴着鸭舌帽，十分淡定，两人快步走进电梯，未接受媒体采访记者了解到，出席高圆圆、赵又廷答谢宴的宾客近百人，其中不少都是女方的高中同学"
for x, w in jieba.analyse.extract_tags(s, withWeight=True):
    print('%s %s' % (x, w))
#此外，还有基于Textrank的关键词抽取办法，最大值为1，从高向低排。


# ## 7-Text Segment and wordcloud

# In[78]:

#define stopword and wordseg
with open("Stopword.txt",encoding="utf-8") as f:
    stopwords = f.read().splitlines()
text = codecs.open('lianghui17.txt', 'r', encoding='utf-8').read() 
seg = [word for word in jieba.cut(text)
                if (word.isalnum())  # Only selecting words that are composed by letters and numbers (exclude marks)
                and (not word.isnumeric())  # Exclude numbers
                and (word not in stopwords)]  # 
text_seg=  ' '.join(seg)
with codecs.open('lianghui17-seg.txt', 'w', 'utf-8') as f:
    f.write(text_seg)


# In[79]:

from wordcloud import WordCloud#部分情况需要使用whl的方式安装
#http://www.lfd.uci.edu/~gohlke/pythonlibs/#wordcloud
wc = WordCloud( background_color = 'white',    # 设置背景颜色
               
                max_words = 300,            # 设置最大现实的字数
                stopwords = stopwords,        # 设置停用词
                font_path = 'msyh.ttf',# 设置字体格式，如不设置显示不了中文
                max_font_size = 300,            # 设置字体最大值
                random_state = 30,            # 设置有多少种随机生成状态，即有多少种配色方案
                width=2000, height=1000)
wc.generate(text_seg)
plt.imshow(wc)
plt.axis('off')
plt.show()
wc.to_file("lianghui-white.png")


# ## 8-Topic model

# In[80]:

report= pd.read_csv('samgov.csv',encoding="gbk")
report.head()


# In[58]:

txt_list = []
# 词性对照表1 http://blog.csdn.net/kevin_darkelf/article/details/39520881
# 词性对照表标准：
cixing=["x","zg","uj","ul","e","d","uz","y","eng","m"]
for i, txt in enumerate(report['content']):
    result = ''
    try:
        for w in pseg.cut(txt):
                if not str(w.flag)  in cixing:
                    seg=str(w.word) ## 不要原来的Encoding
                    if seg not in stopwords:
                            result+=str(seg)+" "
        print (i)
    finally:
        pass
    txt_list.append(result)


# In[59]:

print(txt_list)


# In[81]:

def print_top_words(model, feature_names, n_top_words):
    for topic_idx, topic in enumerate(model.components_):
        print("Topic %d:" % topic_idx)
        print(" ".join([feature_names[i]
                        for i in topic.argsort()[:-n_top_words - 1:-1]]))
    print()


# In[83]:

data_samples=txt_list
with open("Stopword.txt",encoding="utf-8") as f:
    lines = f.read().splitlines()
#自主设计
n_features = 500
n_topics = 10
n_top_words = 50


# In[84]:

from sklearn.feature_extraction.text import TfidfVectorizer,CountVectorizer
from time import time

print("Extracting tf features for LDA...")
tf_vectorizer = CountVectorizer(max_df=0.95, min_df=3, max_features=n_features,
stop_words=lines)
t0 = time()
tf = tf_vectorizer.fit_transform(data_samples)
print("done in %0.3fs." % (time() - t0))


# In[85]:

# Use tf-idf features for NMF.
print("Extracting tf-idf features for NMF...")
tfidf_vectorizer = TfidfVectorizer(max_df=0.95, min_df=3, stop_words=lines)
t0 = time()
tfidf = tfidf_vectorizer.fit_transform(data_samples)
print("done in %0.3fs." % (time() - t0))

# Use tf (raw term count) features for LDA.
print("Extracting tf features for LDA...")
tf_vectorizer = CountVectorizer(max_df=0.95, min_df=3, max_features=n_features,
stop_words=lines)
t0 = time()
tf = tf_vectorizer.fit_transform(data_samples)
print("done in %0.3fs." % (time() - t0))


# In[86]:

# Fit the NMF model
# 事后命名
print("Fitting the NMF model with tf-idf features,"
        "n_features=%d..."
        % (n_features))
t0 = time()
nmf = NMF(n_components=n_topics, random_state=1, alpha=.1, l1_ratio=.5).fit(tfidf)
print("done in %0.3fs." % (time() - t0))

print("Fitting LDA models with tf features, n_features=%d..."
        % (n_features))
lda = LatentDirichletAllocation(n_topics=n_topics, max_iter=5,
learning_method='online', learning_offset=50.,random_state=0)
t0 = time()
lda.fit(tf)
print("done in %0.3fs." % (time() - t0))

print("\nTopics in LDA model")
tf_feature_names = tf_vectorizer.get_feature_names()
print_top_words(lda, tf_feature_names, n_top_words)


# In[ ]:



